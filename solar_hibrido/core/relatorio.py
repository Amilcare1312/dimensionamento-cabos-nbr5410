"""Montagem do memorial de cálculo e exportação do relatório (PDF / DOCX).

`montar_memorial` é uma função pura (testável) que produz uma estrutura de seções
mostrando cada fórmula e valor intermediário. As funções de exportação importam
as bibliotecas de escrita sob demanda (fpdf2 para PDF, python-docx para DOCX).
"""
from __future__ import annotations

from dataclasses import dataclass, field

from .modelos import (
    Carga,
    CriterioFV,
    DadosFatura,
    Irradiacao,
    ParametrosProjeto,
)
from .calculo_baterias import ResultadoBaterias
from .calculo_fv import ResultadoFV
from .calculo_inversor import ResultadoInversor


@dataclass
class Secao:
    titulo: str
    linhas: list[str] = field(default_factory=list)


@dataclass
class Memorial:
    titulo: str
    resumo: list[str] = field(default_factory=list)   # resumo executivo (destaques)
    secoes: list[Secao] = field(default_factory=list)  # memorial detalhado


def _fmt(v: float, casas: int = 1) -> str:
    return f"{v:,.{casas}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def montar_memorial(
    fatura: DadosFatura,
    cargas: list[Carga],
    p: ParametrosProjeto,
    irradiacao: Irradiacao,
    baterias: ResultadoBaterias,
    fv: ResultadoFV,
    inversor: ResultadoInversor,
) -> Memorial:
    """Constrói o memorial de cálculo completo, com fórmulas e valores intermediários."""
    hsp = fv.hsp_utilizado
    criterio_hsp = "mês crítico" if p.criterio_hsp.value == "mes_critico" else "média anual"

    resumo = [
        f"Potência FV: {_fmt(fv.potencia_fv_instalada_kwp, 2)} kWp "
        f"({fv.num_modulos} módulos de {_fmt(p.potencia_modulo_wp, 0)} Wp)",
        f"Banco de baterias: {_fmt(baterias.capacidade_banco_comercial_kwh, 2)} kWh "
        f"({_fmt(baterias.capacidade_nominal_ah, 0)} Ah @ {baterias.tensao_banco_v} V nominais)",
        f"Inversor híbrido: {_fmt(inversor.potencia_nominal_recomendada_kw, 2)} kW "
        f"(surto ≥ {_fmt(inversor.potencia_surto_requerida_kw, 2)} kW)",
        f"Irradiação (HSP {criterio_hsp}): {_fmt(hsp, 2)} kWh/m²/dia — fonte: {irradiacao.fonte}",
    ]

    # ----- Seção 1: dados de entrada -----
    s_entrada = Secao("1. Dados de entrada")
    s_entrada.linhas += [
        f"Distribuidora: {fatura.distribuidora or '—'} | Classe: {fatura.classe_tarifaria} | "
        f"Tensão: {fatura.tensao_fornecimento or '—'}",
        f"Local: {fatura.cidade or '—'}/{fatura.uf or '—'}",
        f"Consumo médio mensal: {_fmt(fatura.consumo_medio_mensal_kwh)} kWh "
        f"({_fmt(fatura.consumo_medio_diario_kwh, 2)} kWh/dia)",
        f"Cargas prioritárias cadastradas: {len(cargas)} "
        f"({sum(1 for c in cargas if c.critica)} críticas)",
    ]

    # ----- Seção 2: irradiação -----
    s_irr = Secao("2. Irradiação solar (HSP)")
    meses = ["Jan", "Fev", "Mar", "Abr", "Mai", "Jun",
             "Jul", "Ago", "Set", "Out", "Nov", "Dez"]
    s_irr.linhas.append(f"Fonte: {irradiacao.fonte}")
    if irradiacao.latitude is not None:
        s_irr.linhas.append(
            f"Coordenadas: {_fmt(irradiacao.latitude, 4)}, {_fmt(irradiacao.longitude, 4)}"
        )
    s_irr.linhas.append(
        "HSP mensal (kWh/m²/dia): "
        + ", ".join(f"{m}={_fmt(v, 2)}" for m, v in zip(meses, irradiacao.hsp_mensal))
    )
    s_irr.linhas.append(
        f"Média anual: {_fmt(irradiacao.hsp_media_anual, 2)} | "
        f"Mês crítico: {_fmt(irradiacao.hsp_mes_critico, 2)} | "
        f"Critério de dimensionamento: {criterio_hsp} ({_fmt(hsp, 2)})"
    )

    # ----- Seção 3: banco de baterias -----
    s_bat = Secao("3. Banco de baterias (backup das cargas críticas)")
    s_bat.linhas += [
        "Fórmula: E_dia = Σ(P × qtd × h) × f_simult",
        f"  E_dia = {_fmt(baterias.energia_diaria_wh, 0)} Wh/dia "
        f"({_fmt(baterias.energia_diaria_kwh, 2)} kWh/dia); f_simult = {p.fator_simultaneidade}",
        "Fórmula: C_util = E_dia × autonomia / η_sistema",
        f"  C_util = {_fmt(baterias.energia_diaria_wh, 0)} × {p.autonomia_dias} / "
        f"{p.eficiencia_sistema} = {_fmt(baterias.capacidade_util_wh, 0)} Wh",
        "Fórmula: C_nom = C_util / DoD",
        f"  C_nom = {_fmt(baterias.capacidade_util_wh, 0)} / {baterias.dod} = "
        f"{_fmt(baterias.capacidade_nominal_wh, 0)} Wh ({_fmt(baterias.capacidade_nominal_kwh, 2)} kWh)",
        f"Conversão para Ah: {_fmt(baterias.capacidade_nominal_wh, 0)} Wh / "
        f"{baterias.tensao_banco_v} V = {_fmt(baterias.capacidade_nominal_ah, 0)} Ah",
        f"Arranjo comercial sugerido: {baterias.num_unidades} unidade(s) de "
        f"{_fmt(baterias.capacidade_unidade_wh/1000, 2)} kWh "
        f"({baterias.unidades_serie} em série × {baterias.unidades_paralelo} em paralelo) "
        f"= {_fmt(baterias.capacidade_banco_comercial_kwh, 2)} kWh instalados",
    ]

    # ----- Seção 4: gerador FV -----
    criterio_fv_txt = (
        "consumo total da fatura"
        if fv.criterio_fv == CriterioFV.CONSUMO_TOTAL
        else "consumo total + recarga do banco"
    )
    s_fv = Secao("4. Gerador fotovoltaico")
    s_fv.linhas += [
        f"Critério: {criterio_fv_txt}",
        f"Consumo diário: {_fmt(fv.consumo_diario_wh, 0)} Wh/dia" +
        (f" + recarga {_fmt(fv.energia_recarga_wh, 0)} Wh" if fv.energia_recarga_wh else ""),
        "Fórmula: E_ger = (consumo + recarga) / η_sistema",
        f"  E_ger = {_fmt(fv.geracao_diaria_necessaria_wh, 0)} Wh/dia",
        "Fórmula: P_fv = E_ger / HSP / PR",
        f"  P_fv = {_fmt(fv.geracao_diaria_necessaria_wh, 0)} / {_fmt(hsp, 2)} / "
        f"{fv.performance_ratio} = {_fmt(fv.potencia_fv_necessaria_wp, 0)} Wp",
        f"Nº de módulos = ⌈P_fv / P_módulo⌉ = ⌈{_fmt(fv.potencia_fv_necessaria_wp, 0)} / "
        f"{_fmt(fv.potencia_modulo_wp, 0)}⌉ = {fv.num_modulos} módulos",
        f"Potência instalada: {_fmt(fv.potencia_fv_instalada_wp, 0)} Wp "
        f"({_fmt(fv.potencia_fv_instalada_kwp, 2)} kWp)",
        f"Área estimada: {_fmt(fv.area_estimada_m2, 1)} m² "
        f"(η_módulo = {int(p.eficiencia_modulo*100)}%)",
        f"Geração diária estimada da usina: {_fmt(fv.geracao_diaria_estimada_wh/1000, 1)} kWh/dia",
    ]

    # ----- Seção 5: inversor -----
    s_inv = Secao("5. Inversor híbrido")
    s_inv.linhas += [
        "Fórmula: P_sim = Σ(P_ativa × qtd) × f_simult",
        f"  P_sim = {_fmt(inversor.potencia_ativa_simultanea_w, 0)} W",
        f"Com margem de {int(inversor.margem*100)}%: "
        f"{_fmt(inversor.potencia_com_margem_w, 0)} W",
        "Pico de partida = Σ(P_regime) + max(surto_i − regime_i)",
        f"  Pico = {_fmt(inversor.pico_partida_w, 0)} W",
        f"Mínimo para acomodar a FV (P_fv / FDI_max = {inversor.faixa_fdi[1]}): "
        f"{_fmt(inversor.potencia_por_fv_w, 0)} W",
        f"Potência nominal recomendada: {_fmt(inversor.potencia_nominal_recomendada_w, 0)} W "
        f"({_fmt(inversor.potencia_nominal_recomendada_kw, 2)} kW)",
        f"Capacidade de surto requerida: ≥ {_fmt(inversor.potencia_surto_requerida_w, 0)} W",
        f"Tensão de entrada do banco compatível: {inversor.tensao_banco_v} V",
    ]

    return Memorial(
        titulo="Memorial de cálculo — Sistema fotovoltaico híbrido",
        resumo=resumo,
        secoes=[s_entrada, s_irr, s_bat, s_fv, s_inv],
    )


def memorial_para_texto(mem: Memorial) -> str:
    """Serializa o memorial em texto simples (útil para testes e pré-visualização)."""
    linhas = [mem.titulo, "=" * len(mem.titulo), "", "RESUMO EXECUTIVO"]
    linhas += [f"  • {r}" for r in mem.resumo]
    linhas.append("")
    for s in mem.secoes:
        linhas.append(s.titulo)
        linhas += [f"    {l}" for l in s.linhas]
        linhas.append("")
    return "\n".join(linhas)


def exportar_pdf(mem: Memorial, caminho: str) -> str:
    """Exporta o memorial para PDF usando fpdf2. Retorna o caminho do arquivo."""
    from fpdf import FPDF  # fpdf2

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def escrever(texto: str, tam: int = 10, estilo: str = "") -> None:
        pdf.set_font("Helvetica", estilo, tam)
        # Latin-1 é o encoding padrão das fontes core do fpdf.
        seguro = texto.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 5, seguro)

    escrever(mem.titulo, 14, "B")
    pdf.ln(2)
    escrever("Resumo executivo", 12, "B")
    for r in mem.resumo:
        escrever(f"- {r}")
    pdf.ln(2)
    for s in mem.secoes:
        escrever(s.titulo, 11, "B")
        for l in s.linhas:
            escrever(l)
        pdf.ln(1)

    pdf.output(caminho)
    return caminho


def exportar_docx(mem: Memorial, caminho: str) -> str:
    """Exporta o memorial para DOCX usando python-docx. Retorna o caminho do arquivo."""
    from docx import Document

    doc = Document()
    doc.add_heading(mem.titulo, level=0)
    doc.add_heading("Resumo executivo", level=1)
    for r in mem.resumo:
        doc.add_paragraph(r, style="List Bullet")
    for s in mem.secoes:
        doc.add_heading(s.titulo, level=1)
        for l in s.linhas:
            doc.add_paragraph(l)
    doc.save(caminho)
    return caminho
